import re
import unicodedata
from pathlib import Path

from docx import Document


class ImportadorIndicadoresMateriasWordError(Exception):
    """Error controlado al analizar la Cédula 4.2.1b."""


PATRON_ATRIBUTO = re.compile(r'\bAE\s*(\d+)\b', re.IGNORECASE)
PATRON_CRITERIO = re.compile(r'^CD\s*(\d+)$', re.IGNORECASE)
PATRON_INDICADOR = re.compile(r'^I\s*(\d+)$', re.IGNORECASE)
PATRON_CLAVE = re.compile(r'^[A-ZÑ]{2,5}-\d{3,5}$')


def limpiar_texto(valor):
    """
    Limpia espacios, saltos de línea y caracteres invisibles provenientes
    del documento Word.
    """
    if valor is None:
        return ''

    texto = str(valor)

    texto = unicodedata.normalize('NFKC', texto)
    texto = texto.replace('\xa0', ' ')
    texto = texto.replace('\u200b', '')
    texto = texto.replace('\ufeff', '')
    texto = texto.replace('\r', ' ')
    texto = texto.replace('\n', ' ')

    return re.sub(r'\s+', ' ', texto).strip()


def normalizar_codigo(valor):
    """
    Normaliza códigos como:

    CD1
    CD 1
    I1
    I 1
    """
    return re.sub(
        r'\s+',
        '',
        limpiar_texto(valor).upper()
    )


def normalizar_clave(valor):
    """
    Normaliza claves como:

    INF - 1028
    INF-1028
    INF – 1028

    a:

    INF-1028
    """
    texto = limpiar_texto(valor).upper()

    texto = re.sub(
        r'[\u2010\u2011\u2012\u2013\u2014\u2015\u2212\u00ad]',
        '-',
        texto
    )

    texto = re.sub(r'\s*-\s*', '-', texto)
    texto = re.sub(r'\s+', '', texto)

    return texto


def extraer_codigo_atributo(documento):
    """
    Busca el código del atributo en las primeras filas del documento.

    Ejemplo:
    AE1. Diseña, mejora e integra sistemas productivos...
    """
    for tabla in documento.tables:
        for fila in tabla.rows[:10]:
            for celda in fila.cells:
                texto = limpiar_texto(celda.text)
                coincidencia = PATRON_ATRIBUTO.search(texto)

                if coincidencia:
                    return f'AE{coincidencia.group(1)}'

    # Como respaldo, buscar en párrafos normales.
    for parrafo in documento.paragraphs:
        texto = limpiar_texto(parrafo.text)
        coincidencia = PATRON_ATRIBUTO.search(texto)

        if coincidencia:
            return f'AE{coincidencia.group(1)}'

    raise ImportadorIndicadoresMateriasWordError(
        'No se pudo identificar el atributo de egreso del documento.'
    )


def encontrar_tabla_asignaciones(documento):
    """
    Encuentra la tabla que contiene las columnas:

    CD
    In
    4.a Clave
    4.b Nombre
    """
    for tabla in documento.tables:
        for fila in tabla.rows:
            textos = [
                limpiar_texto(celda.text).lower()
                for celda in fila.cells
            ]

            fila_completa = ' | '.join(textos)

            if (
                '4.a clave' in fila_completa
                and '4.b nombre' in fila_completa
            ):
                return tabla

    raise ImportadorIndicadoresMateriasWordError(
        'No se encontró la tabla de asignación de indicadores a materias.'
    )


def separar_grupo(valor):
    """
    Convierte valores como:

    7A -> semestre 7, grupo A
    8A -> semestre 8, grupo A
    A  -> semestre None, grupo A

    El resultado se usa únicamente como validación.
    """
    texto = limpiar_texto(valor).upper().replace(' ', '')

    if not texto:
        return {
            'grupo_documento': '',
            'semestre_documento': None,
            'grupo_normalizado': '',
        }

    coincidencia = re.fullmatch(r'(\d+)([A-Z]+)', texto)

    if coincidencia:
        return {
            'grupo_documento': texto,
            'semestre_documento': int(coincidencia.group(1)),
            'grupo_normalizado': coincidencia.group(2),
        }

    return {
        'grupo_documento': texto,
        'semestre_documento': None,
        'grupo_normalizado': texto,
    }


def normalizar_valoracion(valor):
    """
    Normaliza Sí / Si / No.
    """
    texto = limpiar_texto(valor).casefold()

    if texto in {'sí', 'si'}:
        return 'SI'

    if texto == 'no':
        return 'NO'

    return ''


def normalizar_meta(valor):
    """
    Convierte valores como 85%, 85 % o 85.00 a Decimal-compatible string.

    Este valor todavía no se guardará en MateriaIndicador, pero se conserva
    por si más adelante se importa ResultadoIndicador.
    """
    texto = limpiar_texto(valor).replace('%', '').replace(',', '.')

    coincidencia = re.search(r'\d+(?:\.\d+)?', texto)

    if not coincidencia:
        return ''

    return coincidencia.group(0)


def analizar_indicadores_materias_word(ruta_archivo):
    """
    Analiza la sección de asignación de indicadores de la Cédula 4.2.1b.

    Soporta documentos donde Word devuelve 11 o 12 columnas debido
    a celdas combinadas o repetidas.
    """
    ruta = Path(ruta_archivo)

    if not ruta.exists():
        raise ImportadorIndicadoresMateriasWordError(
            'El archivo Word no existe o ya fue eliminado.'
        )

    if ruta.suffix.lower() != '.docx':
        raise ImportadorIndicadoresMateriasWordError(
            'El archivo debe tener formato Word .docx.'
        )

    try:
        documento = Document(str(ruta))
    except Exception as error:
        raise ImportadorIndicadoresMateriasWordError(
            f'No se pudo abrir el documento Word: {error}'
        ) from error

    if not documento.tables:
        raise ImportadorIndicadoresMateriasWordError(
            'El documento no contiene tablas.'
        )

    atributo_codigo = extraer_codigo_atributo(documento)
    tabla = encontrar_tabla_asignaciones(documento)

    asignaciones = []
    filas_omitidas = []
    encabezado_encontrado = False

    for numero_fila, fila in enumerate(tabla.rows, start=1):
        celdas = [
            limpiar_texto(celda.text)
            for celda in fila.cells
        ]

        fila_completa = ' | '.join(celdas).lower()

        # Ignorar filas anteriores al encabezado de la tabla inferior.
        if not encabezado_encontrado:
            if (
                '4.a clave' in fila_completa
                and '4.b nombre' in fila_completa
            ):
                encabezado_encontrado = True

            continue

        # --------------------------------------------------------------
        # Buscar el criterio sin depender de una posición fija.
        # --------------------------------------------------------------
        criterio_codigo = ''

        for valor in celdas:
            codigo = normalizar_codigo(valor)

            if PATRON_CRITERIO.fullmatch(codigo):
                criterio_codigo = codigo
                break

        # --------------------------------------------------------------
        # Buscar el indicador sin depender de una posición fija.
        # En algunos Word aparece repetido: I1, I1.
        # --------------------------------------------------------------
        indicador_codigo = ''

        for valor in celdas:
            codigo = normalizar_codigo(valor)

            if PATRON_INDICADOR.fullmatch(codigo):
                indicador_codigo = codigo
                break

        # --------------------------------------------------------------
        # Buscar dinámicamente la clave de materia.
        # Esto funciona tanto si está en la columna 2 como en la 3.
        # --------------------------------------------------------------
        indice_clave = None
        clave_materia = ''

        for indice, valor in enumerate(celdas):
            posible_clave = normalizar_clave(valor)

            if PATRON_CLAVE.fullmatch(posible_clave):
                indice_clave = indice
                clave_materia = posible_clave
                break

        # Las filas I2/I3/I4 vacías no contienen materia y se ignoran.
        if indice_clave is None:
            continue

        # --------------------------------------------------------------
        # Nombre y grupo aparecen después de la clave.
        # --------------------------------------------------------------
        nombre_materia = ''

        if indice_clave + 1 < len(celdas):
            nombre_materia = limpiar_texto(
                celdas[indice_clave + 1]
            )

        grupo_texto = ''

        if indice_clave + 2 < len(celdas):
            grupo_texto = limpiar_texto(
                celdas[indice_clave + 2]
            )

        # --------------------------------------------------------------
        # Los últimos cuatro valores conservan posiciones estables:
        # periodo, responsable, valoración y meta.
        # --------------------------------------------------------------
        periodo_evaluado = ''
        responsable = ''
        valoracion = ''
        meta = ''

        if len(celdas) >= 4:
            periodo_evaluado = limpiar_texto(celdas[-4])
            responsable = limpiar_texto(celdas[-3])
            valoracion = normalizar_valoracion(celdas[-2])
            meta = normalizar_meta(celdas[-1])

        # --------------------------------------------------------------
        # El instrumento está entre grupo y las últimas cuatro columnas.
        # Puede aparecer duplicado por celdas combinadas.
        # --------------------------------------------------------------
        instrumento = ''

        inicio_instrumento = indice_clave + 3
        fin_instrumento = max(inicio_instrumento, len(celdas) - 4)

        candidatos_instrumento = []

        for valor in celdas[inicio_instrumento:fin_instrumento]:
            texto = limpiar_texto(valor)

            if texto and texto not in candidatos_instrumento:
                candidatos_instrumento.append(texto)

        if candidatos_instrumento:
            instrumento = candidatos_instrumento[0]

        # --------------------------------------------------------------
        # Validaciones.
        # --------------------------------------------------------------
        if not PATRON_CRITERIO.fullmatch(criterio_codigo):
            filas_omitidas.append({
                'fila': numero_fila,
                'motivo': 'No se pudo identificar el criterio.',
                'criterio': criterio_codigo,
                'indicador': indicador_codigo,
                'clave': clave_materia,
            })
            continue

        if not PATRON_INDICADOR.fullmatch(indicador_codigo):
            filas_omitidas.append({
                'fila': numero_fila,
                'motivo': 'No se pudo identificar el indicador.',
                'criterio': criterio_codigo,
                'indicador': indicador_codigo,
                'clave': clave_materia,
            })
            continue

        if not nombre_materia:
            filas_omitidas.append({
                'fila': numero_fila,
                'motivo': 'No se encontró el nombre de la materia.',
                'criterio': criterio_codigo,
                'indicador': indicador_codigo,
                'clave': clave_materia,
            })
            continue

        datos_grupo = separar_grupo(grupo_texto)

        asignaciones.append({
            'fila_documento': numero_fila,

            'atributo_codigo': atributo_codigo,
            'criterio_codigo': criterio_codigo,
            'indicador_codigo': indicador_codigo,

            'clave_materia': clave_materia,
            'nombre_materia': nombre_materia,

            'grupo_documento': datos_grupo['grupo_documento'],
            'semestre_documento': datos_grupo['semestre_documento'],
            'grupo_normalizado': datos_grupo['grupo_normalizado'],

            # Se conservan para una posible importación posterior.
            'instrumento': instrumento,
            'periodo_evaluado': periodo_evaluado,
            'responsable': responsable,
            'valoracion': valoracion,
            'meta': meta,
        })

        # Diagnóstico temporal.
        print(
            'ASIGNACIÓN WORD:',
            atributo_codigo,
            criterio_codigo,
            indicador_codigo,
            clave_materia,
            nombre_materia,
            grupo_texto
        )

    if not encabezado_encontrado:
        raise ImportadorIndicadoresMateriasWordError(
            'No se encontró el encabezado de asignación de materias.'
        )

    if not asignaciones:
        raise ImportadorIndicadoresMateriasWordError(
            'No se encontraron asignaciones válidas de indicadores a materias.'
        )

    return {
        'atributo_codigo': atributo_codigo,
        'asignaciones': asignaciones,
        'filas_omitidas': filas_omitidas,
        'total_asignaciones': len(asignaciones),
        'total_filas_omitidas': len(filas_omitidas),
    }