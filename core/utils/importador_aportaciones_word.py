import re
import unicodedata
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


class ImportadorAportacionesWordError(Exception):
    """Error controlado durante el análisis del mapa de aportación."""


# Claves como INC-1025, ACC-0906, DOC-2307, etc.
PATRON_CLAVE_MATERIA = re.compile(r'^[A-ZÁÉÍÓÚÑ]{2,5}-\d{3,5}$')

NIVELES_VALIDOS = {'I', 'M', 'A'}

DESCRIPCIONES_POR_DEFECTO = {
    'AE1': 'Diseña, mejora e integra sistemas productivos.',
    'AE2': 'Diseña, implementa y mejora sistemas de trabajo.',
    'AE3': 'Implanta sistemas de calidad utilizando métodos estadísticos.',
    'AE4': 'Administra sistemas de mantenimiento.',
    'AE5': 'Gestiona sistemas de seguridad y salud ocupacional.',
    'AE6': 'Formula, evalúa y gestiona proyectos de inversión.',
}


def limpiar_texto(valor):
    """
    Normaliza texto proveniente de una celda Word.
    """
    if valor is None:
        return ''

    texto = str(valor)
    texto = texto.replace('\xa0', ' ')
    texto = texto.replace('\r', ' ')
    texto = texto.replace('\n', ' ')
    texto = re.sub(r'\s+', ' ', texto)

    return texto.strip()


def normalizar_clave(valor):
    """
    Normaliza claves como:
    INC-1009
    INC-1009
    INC – 1009
    INC-1009 con espacios invisibles
    """
    texto = limpiar_texto(valor)

    # Normalizar caracteres Unicode
    texto = unicodedata.normalize('NFKC', texto)

    # Convertir distintos tipos de guion al guion normal
    texto = re.sub(
        r'[\u2010\u2011\u2012\u2013\u2014\u2015\u2212\u00ad]',
        '-',
        texto
    )

    # Quitar espacios normales e invisibles
    texto = texto.replace('\u200b', '')
    texto = texto.replace('\ufeff', '')
    texto = re.sub(r'\s+', '', texto)

    return texto.strip().upper()


def normalizar_nivel(valor):
    """
    Obtiene I, M o A desde una celda.
    Devuelve cadena vacía cuando no existe nivel válido.
    """
    texto = limpiar_texto(valor).upper()

    equivalencias = {
        'I': 'I',
        'INICIAL': 'I',
        'INTRODUCTORIO': 'I',

        'M': 'M',
        'MEDIO': 'M',

        'A': 'A',
        'AVANZADO': 'A',
    }

    return equivalencias.get(texto, '')


def obtener_celdas_segun_cuadricula(fila):
    """
    Convierte una fila Word en una lista respetando la cuadrícula real.

    Si una celda ocupa varias columnas mediante gridSpan, su contenido se
    agrega una vez y las posiciones restantes se rellenan con cadenas vacías.
    """
    valores = []

    for celda_xml in fila._tr.tc_lst:
        textos = celda_xml.xpath('.//w:t/text()')
        texto = limpiar_texto(' '.join(textos))

        propiedades = celda_xml.tcPr
        span = 1

        if propiedades is not None:
            grid_span = propiedades.find(qn('w:gridSpan'))

            if grid_span is not None:
                try:
                    span = int(grid_span.get(qn('w:val')))
                except (TypeError, ValueError):
                    span = 1

        valores.append(texto)

        for _ in range(span - 1):
            valores.append('')

    return valores


def extraer_datos_materia(celdas):
    """
    Estructura fija de las filas de la Cédula 4.2.1a:

    0: número
    1: clave
    2: nombre del curso
    3: AE1
    4: AE2
    5: AE3
    6: AE4
    7: AE5
    8: AE6
    """

    if len(celdas) < 9:
        return None

    numero_texto = limpiar_texto(celdas[0])
    clave = normalizar_clave(celdas[1])
    nombre = limpiar_texto(celdas[2])

    if not numero_texto.isdigit():
        return None

    if not PATRON_CLAVE_MATERIA.fullmatch(clave):
        return None

    if not nombre:
        return None

    niveles = [
        normalizar_nivel(celdas[3]),
        normalizar_nivel(celdas[4]),
        normalizar_nivel(celdas[5]),
        normalizar_nivel(celdas[6]),
        normalizar_nivel(celdas[7]),
        normalizar_nivel(celdas[8]),
    ]

    return {
        'numero': int(numero_texto),
        'clave': clave,
        'nombre': nombre,
        'niveles_lista': niveles,
    }

def obtener_textos_fila(fila):
    """
    Devuelve el contenido normalizado de las celdas de una fila.
    """
    return obtener_celdas_segun_cuadricula(fila)


def encontrar_indice_clave(celdas):
    """
    Busca dentro de una fila la posición de una clave de materia.
    """
    for indice, valor in enumerate(celdas):
        clave = normalizar_clave(valor)

        if PATRON_CLAVE_MATERIA.fullmatch(clave):
            return indice

    return None


def obtener_nombre_materia(celdas, indice_clave):
    """
    Obtiene el nombre de la materia.

    En el formato institucional:
    número | clave | nombre | AE1 | AE2 | AE3 | AE4 | AE5 | AE6

    Por lo tanto, el nombre debe estar inmediatamente después de la clave.
    """
    indice_nombre = indice_clave + 1

    if indice_nombre >= len(celdas):
        return '', None

    nombre = limpiar_texto(celdas[indice_nombre])

    if not nombre:
        return '', None

    return nombre, indice_nombre


def extraer_seis_niveles(celdas, indice_nombre=None):
    """
    Toma siempre las últimas seis celdas como las columnas AE1 a AE6.

    Esto es más confiable para el Word institucional, porque evita que
    las celdas combinadas o repetidas desplacen los niveles.
    """
    if len(celdas) < 6:
        return ['', '', '', '', '', '']

    celdas_atributos = celdas[-6:]

    return [
        normalizar_nivel(valor)
        for valor in celdas_atributos
    ]

def detectar_tabla_principal(documento):
    """
    Localiza la tabla que contiene la mayor cantidad de claves de materia.
    Esto evita depender de document.tables[0].
    """
    mejor_tabla = None
    mayor_numero_materias = 0

    for tabla in documento.tables:
        total_materias = 0

        for fila in tabla.rows:
            celdas = obtener_textos_fila(fila)

            if encontrar_indice_clave(celdas) is not None:
                total_materias += 1

        if total_materias > mayor_numero_materias:
            mejor_tabla = tabla
            mayor_numero_materias = total_materias

    if mejor_tabla is None or mayor_numero_materias == 0:
        raise ImportadorAportacionesWordError(
            'No se encontró la tabla del mapa de aportación dentro del documento.'
        )

    return mejor_tabla


def extraer_descripciones_atributos(tabla, primera_fila_materia):
    """
    Intenta obtener del encabezado las seis descripciones de los atributos.

    Si el formato Word no permite identificarlas correctamente por sus
    celdas combinadas, utiliza las descripciones institucionales por defecto.
    """
    candidatos = []

    for indice_fila in range(primera_fila_materia):
        celdas = obtener_textos_fila(tabla.rows[indice_fila])

        textos_utiles = []

        for valor in celdas:
            valor = limpiar_texto(valor)

            if not valor:
                continue

            if valor.isdigit():
                continue

            if valor.lower() in {
                '#',
                'clave',
                'nombre',
                'nombre del curso',
                '1.a clave',
                '1.b. nombre del curso',
            }:
                continue

            if len(valor) >= 20:
                textos_utiles.append(valor)

        # Quitar duplicados provocados por celdas combinadas.
        textos_unicos = []

        for texto in textos_utiles:
            if texto not in textos_unicos:
                textos_unicos.append(texto)

        if len(textos_unicos) >= 6:
            candidatos = textos_unicos[-6:]

    descripciones = {}

    if len(candidatos) == 6:
        for numero, descripcion in enumerate(candidatos, start=1):
            descripciones[f'AE{numero}'] = descripcion
    else:
        descripciones = DESCRIPCIONES_POR_DEFECTO.copy()

    return descripciones


def analizar_mapa_aportacion_word(ruta_archivo):
    """
    Analiza la Cédula 4.2.1a en formato Word y devuelve:

    {
        'atributos': [
            {
                'codigo': 'AE1',
                'descripcion': '...'
            }
        ],
        'aportaciones': [
            {
                'numero': 1,
                'clave': 'ACC-0906',
                'nombre': 'Fundamentos de Investigación',
                'atributo_codigo': 'AE1',
                'nivel_aporte': 'I'
            }
        ],
        'materias': [
            {
                'numero': 1,
                'clave': 'ACC-0906',
                'nombre': 'Fundamentos de Investigación',
                'niveles': {
                    'AE1': 'I',
                    'AE2': '',
                    'AE3': '',
                    'AE4': '',
                    'AE5': '',
                    'AE6': 'I',
                }
            }
        ],
        'total_materias': 54,
        'total_aportaciones': 120,
    }
    """

    ruta = Path(ruta_archivo)

    if not ruta.exists():
        raise ImportadorAportacionesWordError(
            'El archivo Word no existe o ya fue eliminado.'
        )

    if ruta.suffix.lower() != '.docx':
        raise ImportadorAportacionesWordError(
            'El archivo debe tener formato Word .docx.'
        )

    try:
        documento = Document(str(ruta))
    except Exception as error:
        raise ImportadorAportacionesWordError(
            f'No se pudo abrir el documento Word: {error}'
        ) from error

    if not documento.tables:
        raise ImportadorAportacionesWordError(
            'El documento no contiene tablas.'
        )

    tabla = detectar_tabla_principal(documento)

    filas_materias = []

    for indice_fila, fila in enumerate(tabla.rows):
        celdas = obtener_textos_fila(fila)

        datos = extraer_datos_materia(celdas)

        if not datos:
            continue

        datos['indice_fila'] = indice_fila
        filas_materias.append(datos)

        # Diagnóstico temporal.
        # Puedes quitar este print cuando verifiques que todo coincide.
        print(
            'FILA WORD:',
            datos['clave'],
            '|',
            datos['nombre'],
            '|',
            datos['niveles_lista'],
            '| CELDAS:',
            celdas
        )

    if not filas_materias:
        raise ImportadorAportacionesWordError(
            'No se encontraron materias válidas dentro del documento.'
        )

    primera_fila_materia = min(
        fila['indice_fila']
        for fila in filas_materias
    )

    descripciones = extraer_descripciones_atributos(
        tabla,
        primera_fila_materia
    )

    atributos = []

    for numero in range(1, 7):
        codigo = f'AE{numero}'

        atributos.append({
            'codigo': codigo,
            'descripcion': descripciones.get(
                codigo,
                DESCRIPCIONES_POR_DEFECTO[codigo]
            ),
        })

    materias = []
    aportaciones = []

    for fila in filas_materias:
        niveles_por_atributo = {}

        for numero, nivel in enumerate(
            fila['niveles_lista'],
            start=1
        ):
            codigo_atributo = f'AE{numero}'

            niveles_por_atributo[codigo_atributo] = nivel

            if nivel not in NIVELES_VALIDOS:
                continue

            aportaciones.append({
                'numero': fila['numero'],
                'clave': fila['clave'],
                'nombre': fila['nombre'],
                'atributo_codigo': codigo_atributo,
                'nivel_aporte': nivel,
            })

        materias.append({
            'numero': fila['numero'],
            'clave': fila['clave'],
            'nombre': fila['nombre'],
            'niveles': niveles_por_atributo,
        })

    if not aportaciones:
        raise ImportadorAportacionesWordError(
            'Se encontraron materias, pero no se detectaron niveles I, M o A.'
        )

    return {
        'atributos': atributos,
        'aportaciones': aportaciones,
        'materias': materias,
        'total_materias': len(materias),
        'total_aportaciones': len(aportaciones),
    }